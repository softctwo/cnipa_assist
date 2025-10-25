"""
专利文档解析服务
"""
import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

@dataclass
class PatentDocument:
    """专利文档数据结构"""
    application_number: Optional[str] = None
    application_date: Optional[str] = None
    title: Optional[str] = None
    applicant: Optional[str] = None
    inventor: Optional[str] = None
    agent: Optional[str] = None
    technical_field: Optional[str] = None
    background_art: Optional[str] = None
    invention_content: Optional[str] = None
    claims: List[str] = None
    description: Optional[str] = None
    drawings: List[str] = None
    abstract: Optional[str] = None
    
    def __post_init__(self):
        if self.claims is None:
            self.claims = []
        if self.drawings is None:
            self.drawings = []

class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx', '.txt']
        
    def parse_document(self, file_path: str) -> Tuple[PatentDocument, Dict]:
        """
        解析专利文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            Tuple[PatentDocument, Dict]: 解析结果和元数据
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                return self._parse_word(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._parse_text(file_path)
        except Exception as e:
            raise Exception(f"文档解析失败: {str(e)}")
    
    def _parse_pdf(self, file_path: Path) -> Tuple[PatentDocument, Dict]:
        """解析PDF文档"""
        metadata = {
            "file_size": file_path.stat().st_size,
            "file_type": "PDF",
            "pages": 0,
            "parsing_method": "pdfplumber"
        }
        
        text_content = ""
        
        try:
            # 使用pdfplumber解析PDF
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                        
        except Exception as e:
            # 降级到PyPDF2
            metadata["parsing_method"] = "PyPDF2"
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            except Exception as e2:
                raise Exception(f"PDF解析失败: {str(e2)}")
        
        # 解析文本内容
        patent_doc = self._extract_patent_info(text_content)
        
        return patent_doc, metadata
    
    def _parse_word(self, file_path: Path) -> Tuple[PatentDocument, Dict]:
        """解析Word文档"""
        metadata = {
            "file_size": file_path.stat().st_size,
            "file_type": "Word",
            "parsing_method": "python-docx"
        }
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            metadata["paragraphs"] = len(doc.paragraphs)
            metadata["tables"] = len(doc.tables)
            
            # 解析文本内容
            patent_doc = self._extract_patent_info(text_content)
            
            return patent_doc, metadata
            
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")
    
    def _parse_text(self, file_path: Path) -> Tuple[PatentDocument, Dict]:
        """解析纯文本文档"""
        metadata = {
            "file_size": file_path.stat().st_size,
            "file_type": "Text",
            "parsing_method": "text"
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
        except UnicodeDecodeError:
            # 尝试GBK编码
            with open(file_path, 'r', encoding='gbk') as file:
                text_content = file.read()
                metadata["encoding"] = "gbk"
        
        # 解析文本内容
        patent_doc = self._extract_patent_info(text_content)
        
        return patent_doc, metadata
    
    def _extract_patent_info(self, text: str) -> PatentDocument:
        """从文本中提取专利信息"""
        patent_doc = PatentDocument()
        
        # 申请号提取
        app_number_pattern = r'申请号[：:]\s*(\d{13}|\d{4}\d{8})'
        match = re.search(app_number_pattern, text)
        if match:
            patent_doc.application_number = match.group(1)
        
        # 申请日期提取
        date_pattern = r'申请日[：:]\s*(\d{4}[年.-]\d{1,2}[月.-]\d{1,2})'
        match = re.search(date_pattern, text)
        if match:
            patent_doc.application_date = match.group(1)
        
        # 发明名称提取
        title_patterns = [
            r'发明名称[：:]\s*(.+?)(?:\n|申请人)',
            r'实用新型名称[：:]\s*(.+?)(?:\n|申请人)',
            r'名\s*称[：:]\s*(.+?)(?:\n|申请人)'
        ]
        for pattern in title_patterns:
            match = re.search(pattern, text)
            if match:
                patent_doc.title = match.group(1).strip()
                break
        
        # 申请人提取
        applicant_pattern = r'申请人[：:]\s*(.+?)(?:\n|发明人|地址)'
        match = re.search(applicant_pattern, text)
        if match:
            patent_doc.applicant = match.group(1).strip()
        
        # 发明人提取
        inventor_pattern = r'发明人[：:]\s*(.+?)(?:\n|申请人|地址)'
        match = re.search(inventor_pattern, text)
        if match:
            patent_doc.inventor = match.group(1).strip()
        
        # 技术领域提取
        field_patterns = [
            r'技术领域\s*(.+?)(?=背景技术|发明内容|\n\s*\n)',
            r'所属技术领域\s*(.+?)(?=背景技术|发明内容|\n\s*\n)'
        ]
        for pattern in field_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                patent_doc.technical_field = match.group(1).strip()
                break
        
        # 背景技术提取
        bg_pattern = r'背景技术\s*(.+?)(?=发明内容|技术方案|\n\s*\n)'
        match = re.search(bg_pattern, text, re.DOTALL)
        if match:
            patent_doc.background_art = match.group(1).strip()
        
        # 发明内容提取
        content_pattern = r'发明内容\s*(.+?)(?=具体实施方式|附图说明|\n\s*\n)'
        match = re.search(content_pattern, text, re.DOTALL)
        if match:
            patent_doc.invention_content = match.group(1).strip()
        
        # 权利要求提取
        claims_pattern = r'权利要求书\s*(.+?)(?=说明书|附图说明|$)'
        match = re.search(claims_pattern, text, re.DOTALL)
        if match:
            claims_text = match.group(1)
            # 分离各项权利要求
            claim_items = re.findall(r'(\d+)\.\s*(.+?)(?=\d+\.|$)', claims_text, re.DOTALL)
            patent_doc.claims = [f"{num}. {content.strip()}" for num, content in claim_items]
        
        # 摘要提取
        abstract_pattern = r'摘\s*要\s*(.+?)(?=附图说明|权利要求|$)'
        match = re.search(abstract_pattern, text, re.DOTALL)
        if match:
            patent_doc.abstract = match.group(1).strip()
        
        return patent_doc
    
    def validate_document(self, patent_doc: PatentDocument) -> Dict[str, List[str]]:
        """
        验证文档完整性
        
        Args:
            patent_doc: 专利文档对象
            
        Returns:
            Dict[str, List[str]]: 验证结果，包含错误和警告
        """
        errors = []
        warnings = []
        
        # 必填字段检查
        if not patent_doc.title:
            errors.append("缺少发明名称")
        
        if not patent_doc.applicant:
            errors.append("缺少申请人信息")
        
        if not patent_doc.claims:
            errors.append("缺少权利要求书")
        
        if not patent_doc.abstract:
            warnings.append("缺少摘要")
        
        # 格式检查
        if patent_doc.application_number:
            if not re.match(r'^\d{13}$|^\d{12}$', patent_doc.application_number):
                warnings.append("申请号格式可能不正确")
        
        if patent_doc.application_date:
            if not re.match(r'\d{4}[年.-]\d{1,2}[月.-]\d{1,2}', patent_doc.application_date):
                warnings.append("申请日期格式可能不正确")
        
        # 内容完整性检查
        if patent_doc.claims and len(patent_doc.claims) == 0:
            errors.append("权利要求书为空")
        
        if patent_doc.title and len(patent_doc.title) > 25:
            warnings.append("发明名称过长（建议不超过25字）")
        
        return {
            "errors": errors,
            "warnings": warnings
        }

# 创建全局解析器实例
document_parser = DocumentParser()